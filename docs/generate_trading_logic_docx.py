"""
Generate Trading Logic & Edge Framework .docx in EverestQuant IC pack style.

Produces:
  - docs/Trading_Logic_and_Edge_Framework.docx

Style: Garamond fonts, navy headers (#1A3C5E), alternating table shading.
Matches ME Equity L/S IC Pack formatting.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# -----------------------------------------------------------------------
# Paths
# -----------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "Trading_Logic_and_Edge_Framework.docx"

# -----------------------------------------------------------------------
# IC Pack styling constants
# -----------------------------------------------------------------------
GARAMOND = "Garamond"
NAVY = "#1a5276"
DARK_BLUE = "#1a3c5e"
HEADER_ROW_BG = "1A3C5E"
ALT_ROW_BG = "F2F7FB"
LIGHT_GRAY = "F8F9FA"


# -----------------------------------------------------------------------
# Docx helpers (IC pack style)
# -----------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, align: str = "left",
                  font_size: int = 9, color: RGBColor = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = GARAMOND
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_table(doc, headers: list, rows: list, col_aligns: list = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_ROW_BG)
        set_cell_text(cell, h, bold=True, align="center", font_size=9,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            if i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            align = "left"
            if col_aligns and j < len(col_aligns):
                align = col_aligns[j]
            bold = (j == 0)
            set_cell_text(cell, str(val), bold=bold, align=align, font_size=9)
    doc.add_paragraph("")


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = GARAMOND
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)


def add_body(doc, text: str, bold: bool = False, size: int = 11,
             italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = GARAMOND
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text: str, bold_prefix: str = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = GARAMOND
        r.font.size = Pt(11)
        r.bold = True
        r2 = p.add_run(text)
        r2.font.name = GARAMOND
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = GARAMOND
        run.font.size = Pt(11)


def add_formula(doc, text: str) -> None:
    """Add a formula block (indented, monospace-ish)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def add_spacer(doc):
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# -----------------------------------------------------------------------
# Main document generation
# -----------------------------------------------------------------------
def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title Page ──────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Everest Agentic AI Trader")
    run.font.name = GARAMOND
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Trading Logic & Mathematical Edge Framework")
    run.font.name = GARAMOND
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("EverestQuant Research Paper  |  March 2026\nClassification: Internal")
    run.font.name = GARAMOND
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.italic = True

    doc.add_page_break()

    # ── Table of Contents ───────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Why the Edge Exists",
        "3. Fair Probability Estimation",
        "4. Edge Calculation & Confidence Scoring",
        "5. Position Sizing: Fractional Kelly Criterion",
        "6. Merge Arbitrage",
        "7. Learning Agent",
        "8. Statistical Edge Validation",
        "9. Live Performance Test Metrics",
        "10. Summary of Edge Sources",
        "11. Market Matching Architecture",
        "12. Paper Trading Results (March 2026)",
        "13. Testing & Validation Framework",
        "References",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = GARAMOND
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    doc.add_page_break()

    # ── 1. Executive Summary ────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)

    add_body(doc, (
        "The Everest Agentic AI Trader exploits structural inefficiencies between "
        "Polymarket's prediction market CLOB and sharp bookmaker lines (primarily Pinnacle). "
        "The system identifies mispriced outcomes, sizes positions using fractional Kelly "
        "criterion, and learns from resolved trades to continuously improve edge estimation."
    ))
    add_body(doc, (
        "Core thesis: Polymarket sports markets are systematically less efficient than "
        "Pinnacle because (1) they attract recreational flow, (2) they lack professional "
        "market-making infrastructure, and (3) structural features like the 3-second TAKER "
        "delay and neg-risk collateral requirements create friction that prevents instantaneous "
        "arbitrage."
    ))
    add_body(doc, (
        "The system targets 3-15% edge per trade in the 5-95c price range, uses fractional Kelly "
        "sizing ($500 test bankroll, $8 max position), and holds positions to resolution. "
        "Six independent edge sources provide robustness against any single source diminishing."
    ))

    # Key metrics box
    add_heading(doc, "Key Parameters", level=2)
    add_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["Bankroll", "$500", "Test wallet on Polygon"],
            ["Max Position", "$8 (1.6%)", "Small size for statistical learning phase"],
            ["Max Exposure", "$400 (80%)", "Maximise capital deployment"],
            ["Kelly Fraction", "0.15 (fractional)", "Conservative: 15% of full Kelly"],
            ["Min Edge", "3%", "Below this, transaction costs dominate"],
            ["Max Edge", "20%", "Above this, likely matching error"],
            ["Price Range", "5c - 95c", "Full range for h2h/spread/total markets"],
            ["Max Time to Event", "10 days", "Balance capital lockup vs coverage"],
            ["Min Liquidity", "$100", "Ensures reliable midpoint pricing"],
            ["Min 24h Volume", "$1,000", "Avoid moving illiquid markets; prefer high volume"],
            ["Scan Interval", "300 seconds", "Balance API cost vs opportunity capture"],
            ["Target Sports", "26", "EPL, NBA, NFL, NHL, CBB, ATP, WTA, etc."],
        ],
        col_aligns=["left", "center", "left"],
    )

    # ── 2. Why the Edge Exists ──────────────────────────────────────
    add_heading(doc, "2. Why the Edge Exists", level=1)

    add_heading(doc, "2.1 Market Structure Asymmetry", level=2)
    add_body(doc, (
        "Polymarket operates as a binary options CLOB on Polygon. Each outcome trades as a "
        "token priced $0-$1.00, with YES + NO = $1.00 guaranteed by the CTF contract. "
        "Key structural features create exploitable inefficiency:"
    ))
    add_table(doc,
        ["Feature", "Polymarket", "Pinnacle"],
        [
            ["Participant base", "Retail-heavy, crypto-native", "Professional bettors, syndicates"],
            ["Market-making", "Fragmented, often manual", "Algorithmic, sub-second"],
            ["Vig / spread", "1-5% per side", "1.5-3% total overround"],
            ["TAKER delay", "3 seconds (sports)", "None"],
            ["Liquidity", "$10K-$500K per market", "$50K-$2M per market"],
        ],
        col_aligns=["left", "center", "center"],
    )

    add_heading(doc, "2.2 The Favourite-Longshot Bias", level=2)
    add_body(doc, (
        "Prediction markets exhibit the well-documented FLB: longshots are overpriced "
        "relative to their true probability, and favourites are underpriced. On Polymarket, "
        "this manifests as tokens across the full price range trading above or below fair value. "
        "RN1 (the reference trader, +$20.35M verified P&L over 1.1M trades) trades the full "
        "0-100c range (median entry 0.48), suggesting mispricing exists at all price levels, "
        "not just in longshots."
    ))

    add_heading(doc, "2.3 Information Propagation Delay", level=2)
    add_body(doc, (
        "When a material event occurs (injury, lineup change), Pinnacle's lines adjust "
        "within seconds via algorithmic market makers. Polymarket's prices adjust over minutes "
        "due to fewer active makers, lower capital deployed, and the 3-second TAKER delay. "
        "This creates a latency arbitrage window of 30-180 seconds that systematic traders exploit."
    ))

    # ── 3. Fair Probability Estimation ──────────────────────────────
    add_heading(doc, "3. Fair Probability Estimation", level=1)

    add_body(doc, (
        "Bookmaker odds contain embedded margin ('overround'). To extract fair probabilities, "
        "we must remove this margin. The bot uses sport-specific methods because margin "
        "distribution varies by market structure."
    ))

    add_heading(doc, "3.1 Shin's Model (Soccer, 3-way)", level=2)
    add_body(doc, (
        "Shin (1991, 1993) models a bookmaker protecting against insider traders. "
        "Margin is distributed non-equally: longshots carry more margin due to adverse selection."
    ))
    add_formula(doc, "fair_prob_i = [sqrt(z^2 + 4(1-z) * ip_i^2 / S) - z] / [2(1-z)]")
    add_body(doc, (
        "where S = sum of implied probabilities, ip_i = 1/odds_i, and z (insider fraction) "
        "is solved via bisection. Used for: EPL, Bundesliga, La Liga, UCL, Serie A, Ligue 1."
    ))

    add_heading(doc, "3.2 MWPO (US Sports, 2-way)", level=2)
    add_body(doc, (
        "For tight 2-way markets (NBA, NFL), margin is subtracted proportional to decimal odds:"
    ))
    add_formula(doc, "fair_prob_i = ip_i - w_i * M")
    add_formula(doc, "where M = sum(ip) - 1,  w_i = odds_i / sum(odds_j)")
    add_body(doc, "Used for: NBA, NFL, CBB, NHL.")

    add_heading(doc, "3.3 Power Method (Tennis)", level=2)
    add_body(doc, "Each implied probability raised to exponent k, solved via bisection:")
    add_formula(doc, "fair_prob_i = ip_i^k,  where sum(ip_i^k) = 1")
    add_body(doc, "Used for: ATP, WTA tennis (2-way markets with significant FLB at extremes).")

    add_heading(doc, "3.4 Multi-Book Consensus", level=2)
    add_body(doc, "When multiple sharp books are available, fair probabilities are combined via efficiency-weighted average:")
    add_formula(doc, "P_fair = sum(P_fair_j * w_j) / sum(w_j)")
    add_table(doc,
        ["Bookmaker", "Efficiency Weight", "Rationale"],
        [
            ["Pinnacle", "1.00", "Sharpest, tightest margins, largest limits"],
            ["Betfair Exchange", "0.95", "True exchange, but liquidity varies"],
            ["Matchbook", "0.85", "Commission-based exchange"],
            ["BetCris", "0.75", "Sharp Central American book"],
        ],
        col_aligns=["left", "center", "left"],
    )

    # ── 4. Edge Calculation ─────────────────────────────────────────
    add_heading(doc, "4. Edge Calculation & Confidence Scoring", level=1)

    add_heading(doc, "4.1 Raw Edge", level=2)
    add_formula(doc, "edge = P_fair - P_polymarket")
    add_formula(doc, "edge_pct = 100 * edge / P_polymarket")
    add_body(doc, (
        "The edge represents the percentage by which Polymarket underprices an outcome "
        "relative to sharp-book fair value. Filters: min 3%, max 20%, price 5-95c, "
        "exact line match for spreads/totals (0.01 tolerance), max 10 days to event, "
        "min $100 liquidity, min $1,000 24h volume. Markets are sorted by edge first, "
        "then by volume (higher volume preferred) to prioritise liquid markets. "
        "Exotic sub-markets (first-set, set-handicap, corners, etc.) "
        "are excluded to prevent phantom edges from cross-market-type comparison."
    ))

    add_heading(doc, "4.2 Confidence Scoring", level=2)
    add_body(doc, (
        "Not all edges are equally reliable. Five independent factors, each mapped [0,1], "
        "produce a weighted confidence score:"
    ))
    add_table(doc,
        ["Factor", "Weight", "Calculation"],
        [
            ["Book agreement", "0.30", "min(1.0, agreeing_books / total * 1.2)"],
            ["Time to event", "0.25", "1.0 (<2h), 0.85 (<6h), 0.65 (<24h), 0.45 (<48h), 0.30"],
            ["Liquidity", "0.15", "min(1.0, log(1+liq) / log(101000))"],
            ["Market type", "0.15", "h2h: 0.90, spread: 0.75, total: 0.60"],
            ["Historical accuracy", "0.15", "Learning agent win rate for segment"],
        ],
        col_aligns=["left", "center", "left"],
    )

    add_heading(doc, "4.3 Edge Decay Model", level=2)
    add_body(doc, "Edges decay as game start approaches and prices converge to fair value:")
    add_formula(doc, "decay = 0.40 + 0.60 * exp(-hours_to_start / 12)")
    add_formula(doc, "effective_edge = raw_edge * decay")
    add_body(doc, "No decay applied within 2 hours of start (edges are most reliable near close).")

    # ── 5. Kelly Criterion ──────────────────────────────────────────
    add_heading(doc, "5. Position Sizing: Fractional Kelly Criterion", level=1)

    add_heading(doc, "5.1 Full Kelly Formula", level=2)
    add_body(doc, "For a binary bet with net odds b = (1/price) - 1:")
    add_formula(doc, "f* = (b*p - q) / b")
    add_body(doc, "where p = fair probability, q = 1-p. This maximises the expected geometric growth rate:")
    add_formula(doc, "G(f) = p * log(1 + f*b) + q * log(1 - f)")

    add_heading(doc, "5.2 Fractional Kelly Rationale", level=2)
    add_body(doc, (
        "Full Kelly is optimal only with perfect probability estimates and infinite horizon. "
        "In practice, estimation error makes it dangerously aggressive. We use 15% Kelly "
        "(f = 0.15 * f*) during the paper trading / statistical learning phase."
    ))
    add_table(doc,
        ["Kelly Fraction", "Growth Captured", "Variance Captured", "Ruin Probability"],
        [
            ["Full (1.00)", "100%", "100%", "Material"],
            ["Half (0.50)", "75%", "50%", "Low"],
            ["Quarter (0.25)", "56%", "25%", "Negligible"],
            ["15% (0.15) [current]", "~40%", "~15%", "Near zero"],
        ],
        col_aligns=["left", "center", "center", "center"],
    )
    add_body(doc, (
        "At 15% Kelly, position sizes are very small ($0.50-$8.00) which prioritises "
        "maximising trade count for statistical learning over capital growth. Once edge is "
        "confirmed (M2 milestone, n>=100), Kelly fraction can be increased to 0.25."
    ))

    add_heading(doc, "5.3 Estimation-Error Adjustment (Thorp 2006)", level=2)
    add_body(doc, "When edge uncertainty is quantifiable:")
    add_formula(doc, "f_adjusted = f_kelly * [1 - 0.5 * (sigma_edge / edge)]")
    add_body(doc, "This reduces position size when the edge estimate has high variance.")

    add_heading(doc, "5.4 Position Constraints", level=2)
    add_formula(doc, "size = max($0.50, min(f_used * bankroll, $8))")
    add_bullet(doc, " $400 = 80% of $500 bankroll", bold_prefix="Max exposure:")
    add_bullet(doc, " $0.50 minimum (fractional Kelly floor)", bold_prefix="Min trade:")
    add_bullet(doc, " No opposing sides of same event (over+under, yes+no)", bold_prefix="Conflict check:")
    add_bullet(doc, " No duplicate positions on same token", bold_prefix="Duplicate check:")
    add_bullet(doc, " effective_bankroll = initial + realized_pnl", bold_prefix="Dynamic bankroll:")

    add_heading(doc, "5.5 Slippage Estimation & Tracking", level=2)
    add_body(doc, (
        "Before each order, the system fetches the current best ask price from the Polymarket "
        "orderbook and computes expected slippage in basis points:"
    ))
    add_formula(doc, "slippage_bps = (best_ask - expected_price) / expected_price * 10,000")
    add_body(doc, (
        "In live mode, the actual fill price is captured from the order result and compared "
        "to the expected price to compute realised slippage. Both pre-trade and post-trade "
        "slippage are logged per trade for ongoing monitoring. Markets with <$1,000 24h volume "
        "are excluded entirely to minimise market impact."
    ))

    # ── 6. Merge Arbitrage ──────────────────────────────────────────
    add_heading(doc, "6. Merge Arbitrage (Risk-Free)", level=1)

    add_body(doc, (
        "On Polymarket, YES + NO tokens for the same market can be merged into $1.00 "
        "via the CTF contract. When YES_ask + NO_ask < $1.00, a risk-free profit exists:"
    ))
    add_formula(doc, "profit_per_pair = $1.00 - (yes_ask + no_ask)")
    add_formula(doc, "pairs = min(yes_depth, no_depth)")
    add_formula(doc, "total_profit = profit_per_pair * pairs - gas_cost")

    add_body(doc, (
        "Merge opportunities persist because: (1) neg-risk collateral reduces liquidity, "
        "(2) YES/NO market makers operate independently, (3) during volatile events one side "
        "reprices faster. RN1 data shows MERGE was the primary profit mechanism: $40.4M in "
        "synthetic sells (35% of total volume)."
    ))

    # ── 7. Learning Agent ───────────────────────────────────────────
    add_heading(doc, "7. Learning Agent", level=1)

    add_heading(doc, "7.1 Adaptive Edge Adjustment", level=2)
    add_body(doc, (
        "The learning agent segments trades by sport, market type, and price bucket. "
        "With sufficient samples (n >= 20), it adjusts edge estimates:"
    ))
    add_formula(doc, "adjustment = (actual_wr - predicted_wr) / predicted_wr")
    add_formula(doc, "adjusted_edge = raw_edge * (1 + adjustment * 0.30)")
    add_body(doc, "Learning rate of 0.30 is conservative to avoid overfitting to small samples.")

    add_heading(doc, "7.2 Sport Scoring (Capital Allocation)", level=2)
    add_formula(doc, "score = win_rate * log2(n + 1) * (1 + avg_edge / 100)")
    add_body(doc, "Higher scores receive preferential capital allocation across the portfolio.")

    # ── 8. Statistical Validation ───────────────────────────────────
    add_heading(doc, "8. Statistical Edge Validation", level=1)

    add_body(doc, (
        "Before trusting any observed edge, three independent statistical tests are applied:"
    ))

    add_heading(doc, "8.1 Binomial Test (Win Rate)", level=2)
    add_body(doc, "H0: actual_win_rate = market-implied win rate.")
    add_formula(doc, "z = (actual_wr - expected_wr) / sqrt(expected_wr * (1-expected_wr) / n)")
    add_body(doc, "Reject H0 at p < 0.05: win rate is significantly better than market-implied.")

    add_heading(doc, "8.2 t-Test on P&L", level=2)
    add_body(doc, "H0: mean(pnl) = 0.")
    add_formula(doc, "t = mean(pnl) / (stdev(pnl) / sqrt(n))")
    add_body(doc, "Reject H0 at p < 0.05: strategy generates statistically significant positive P&L.")

    add_heading(doc, "8.3 Runs Test (Randomness)", level=2)
    add_body(doc, (
        "Verifies wins/losses are not serially correlated. p > 0.05 (desired): "
        "outcomes appear random, edge is not driven by streaks."
    ))

    add_heading(doc, "8.4 Combined Confidence Score", level=2)
    add_table(doc,
        ["Component", "Points", "Criterion"],
        [
            ["Binomial test", "25", "p < 0.01"],
            ["t-Test", "25", "p < 0.01"],
            ["Sample size", "15", ">= required minimum"],
            ["Flat-bet ROI", "15", "> 5%"],
            ["Runs test", "10", "p > 0.05 (random)"],
            ["Win rate margin", "10", "> expected + 5pp"],
            ["Total", "100", ""],
        ],
        col_aligns=["left", "center", "left"],
    )
    add_body(doc, "Verdict: >= 80 Strong evidence | >= 60 Moderate | >= 40 Inconclusive | < 40 No evidence.", italic=True)

    # ── 9. Live Performance Metrics ─────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "9. Live Performance Test Metrics", level=1)

    add_body(doc, (
        "The following metrics define the scorecard against which live agent trading "
        "performance should be evaluated. Thresholds are calibrated for a $500 test wallet "
        "with 15% fractional Kelly sizing."
    ))

    add_heading(doc, "9.1 Primary Metrics (Must-Pass)", level=2)
    add_table(doc,
        ["Metric", "Formula", "Target", "Red Flag", "Review"],
        [
            ["Closing Line Value", "(entry_fair - closing_fair) / closing_fair", "> +2%", "< 0%", "Weekly"],
            ["Win Rate vs Expected", "actual_wr - mean(entry_prices)", "> +3pp", "< 0pp", "Weekly (n>=30)"],
            ["Flat-Bet ROI", "total_pnl / capital_deployed", "> +3%", "< -5%", "Weekly"],
            ["Binomial p-value", "See Section 8.1", "< 0.10", "> 0.50", "After 50 trades"],
            ["t-Test p-value", "See Section 8.2", "< 0.10", "> 0.50", "After 50 trades"],
        ],
        col_aligns=["left", "left", "center", "center", "center"],
    )

    add_body(doc, (
        "CLV is the single most important metric. A positive CLV means the agent consistently "
        "buys at prices better than where the market closes -- the gold standard of sharp "
        "betting. Even during losing streaks, positive CLV confirms the edge is real."
    ), bold=True)

    add_heading(doc, "9.2 Risk Metrics (Guardrails)", level=2)
    add_table(doc,
        ["Metric", "Target", "Hard Limit", "Action if Breached"],
        [
            ["Max Drawdown", "< 15% bankroll", "25% ($125)", "Halt, review"],
            ["Drawdown Duration", "< 14 days", "30 days", "Reduce sizes 50%"],
            ["Daily Loss Limit", "> -$25", "-$50", "Halt for day"],
            ["Exposure / Bankroll", "< 80%", "90%", "No new positions"],
            ["Single Position / Bankroll", "< 1.6% ($8)", "3% ($15)", "Reject trade"],
            ["Contradictory Positions", "0", "0", "Immediate investigation"],
        ],
        col_aligns=["left", "center", "center", "left"],
    )

    add_heading(doc, "9.3 Edge Quality Metrics (Diagnostic)", level=2)
    add_table(doc,
        ["Metric", "Formula", "Healthy Range", "Concern"],
        [
            ["Brier Score", "mean((pred - actual)^2)", "< 0.22", "> 0.25"],
            ["Log Loss", "-mean(y*log(p) + (1-y)*log(1-p))", "< 0.65", "> 0.69"],
            ["Profit Factor", "gross_wins / |gross_losses|", "> 1.3", "< 1.0"],
            ["Avg Edge Realized", "mean(pnl / cost)", "> +3%", "< 0%"],
            ["Edge Decay Ratio", "realized_edge / entry_edge", "> 0.50", "< 0.30"],
            ["Match Accuracy", "% correct matches", "> 90%", "< 80%"],
        ],
        col_aligns=["left", "left", "center", "center"],
    )

    add_heading(doc, "9.4 Operational Health", level=2)
    add_table(doc,
        ["Metric", "Target", "Red Flag"],
        [
            ["API uptime", "> 99%", "< 95%"],
            ["Scan cycle success rate", "> 98%", "< 90%"],
            ["Odds API requests remaining", "> 5,000 / month", "< 1,000"],
            ["Position resolution latency", "< 24h after event", "> 48h"],
            ["Stale positions (past events)", "0", "> 5"],
            ["404 sport key errors / cycle", "0", "> 2"],
        ],
        col_aligns=["left", "center", "center"],
    )

    add_heading(doc, "9.5 Milestone Checkpoints", level=2)
    add_table(doc,
        ["Milestone", "Trades", "Key Decision"],
        [
            ["M1: Signal Validation", "30", "Is CLV positive? Win rate > expected? If no, review matching."],
            ["M2: Statistical Significance", "100", "Binomial p<0.10? t-test p<0.10? If no, edge may not be real."],
            ["M3: Sizing Validation", "100", "Is Kelly profitable? Compare vs flat-bet. If flat wins, recalibrate."],
            ["M4: Sport Segmentation", "200", "Which sports profitable? Prune losers, concentrate winners."],
            ["M5: Live Readiness", "300", "CLV>+2%, Brier<0.22, PF>1.3. Approve paper-to-live transition."],
        ],
        col_aligns=["left", "center", "left"],
    )

    add_heading(doc, "9.6 Benchmark Comparison", level=2)
    add_table(doc,
        ["Benchmark", "Expected Performance", "Source"],
        [
            ["Random betting (no edge)", "-2% to -5% ROI", "Theoretical (bookmaker vig)"],
            ["Naive Pinnacle follower", "+0% to +2% ROI", "Industry consensus"],
            ["Competent sports bettor", "+3% to +8% CLV", "Pinnacle data"],
            ["RN1 (reference trader)", "+21.8% ROI on $93.1M", "Verified activity data"],
            ["Our target (paper phase)", "+3% CLV, +5% flat-bet ROI", "Conservative for $500 wallet"],
        ],
        col_aligns=["left", "center", "left"],
    )

    # ── 10. Edge Sources Summary ────────────────────────────────────
    add_heading(doc, "10. Summary of Edge Sources", level=1)
    add_table(doc,
        ["Edge Source", "Mechanism", "Magnitude", "Persistence"],
        [
            ["Sharp-book mispricing", "Polymarket lags Pinnacle 30-180s", "3-15%", "High (structural)"],
            ["Favourite-longshot bias", "Longshots overpriced in prediction mkts", "3-20%", "High (behavioural)"],
            ["Merge arbitrage", "YES+NO < $1.00, fragmented liquidity", "1-5%", "High (structural)"],
            ["TAKER delay exploitation", "3s delay prevents rapid correction", "2-8%", "High (protocol)"],
            ["Information asymmetry", "Sharp books price news faster", "Variable", "Medium (competition)"],
            ["Edge decay capture", "Enter early, CLV confirms edge", "Captured via CLV", "Medium (time-dependent)"],
        ],
        col_aligns=["left", "left", "center", "center"],
    )
    add_body(doc, (
        "The combination of multiple independent edge sources provides robustness: "
        "even if one source diminishes (e.g., more market makers enter Polymarket), "
        "others persist due to structural protocol constraints."
    ))

    # ── 11. Market Matching Architecture ─────────────────────────────
    doc.add_page_break()
    add_heading(doc, "11. Market Matching Architecture", level=1)

    add_body(doc, (
        "The matcher is the critical pipeline stage that links Polymarket slug-based markets "
        "to structured odds from The Odds API. Coverage directly determines the opportunity set."
    ))

    add_heading(doc, "11.1 Matching Pipeline", level=2)
    add_body(doc, (
        "For each Polymarket market, the matcher: (1) extracts sport, teams, and date from the "
        "slug, (2) classifies market type (h2h, spread, total, exotic), (3) finds candidate "
        "Odds API events within a 96-hour date window, (4) performs strict team name matching, "
        "(5) for spreads/totals, requires exact line match (0.01 tolerance)."
    ))

    add_heading(doc, "11.2 Team Name Resolution", level=2)
    add_table(doc,
        ["Method", "Scope", "Example"],
        [
            ["TEAM_ALIASES dict", "136 teams (NBA, NFL, EPL, etc.)", "lakers -> Los Angeles Lakers"],
            ["Sport-scoped abbrevs", "32 NHL + 20 CBB teams", "(nhl, edm) -> Edmonton Oilers"],
            ["Tennis surname match", "All ATP/WTA dynamically", "slug 'djokovic' -> 'Novak Djokovic'"],
            ["Fuzzy matching (0.80)", "Fallback for partial matches", "man-utd -> Manchester United"],
        ],
        col_aligns=["left", "left", "left"],
    )

    add_heading(doc, "11.3 Exotic Market Filtering", level=2)
    add_body(doc, (
        "Sub-markets like first-set-winner, set-handicap, set-totals, match-total, corners, "
        "and tiebreak markets are filtered BEFORE market type classification. Without this, "
        "their prices get compared to match-winner odds, creating phantom edges of 10-65%. "
        "This was a critical bug fix: the exotic check must run before spread/total detection "
        "because patterns like '-corners-over-' contain '-over-' which triggers total classification."
    ))

    add_heading(doc, "11.4 Coverage Metrics", level=2)
    add_table(doc,
        ["Metric", "Before (Mar 5)", "After (Mar 6)", "Improvement"],
        [
            ["Configured sports", "22", "26", "+4 (NHL, CBB, CFB, Eredivisie)"],
            ["Team aliases", "136", "188+", "+52 NHL/CBB + dynamic tennis"],
            ["Date match window", "48 hours", "96 hours", "2x wider"],
            ["Polymarket fetch limit", "100/sport", "200/sport", "2x coverage"],
            ["Matched markets/scan", "44", "96", "+118%"],
            ["Trades placed/scan", "0-13", "43-87", "6-7x increase"],
        ],
        col_aligns=["left", "center", "center", "center"],
    )

    # ── 12. Paper Trading Results ──────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "12. Paper Trading Results (March 2026)", level=1)

    add_heading(doc, "12.1 Phase 1: Initial Deployment (March 5)", level=2)
    add_body(doc, (
        "Initial deployment with 22 sports, 136 team aliases, 48h date window. "
        "Matched 44 markets per scan, placed 62 positions over ~12 hours."
    ))
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Positions opened", "62"],
            ["Resolved (NBA)", "15 (7W / 8L)"],
            ["Realized P&L", "-$15.34"],
            ["Win rate", "46.7% (expected ~50%)"],
            ["Avg position size", "$4.50"],
        ],
        col_aligns=["left", "center"],
    )
    add_body(doc, (
        "Key finding: contradictory position bug caused guaranteed losses (bought both Over+Under "
        "on same game totals). Four contradictory pairs identified, contributing -$14 to losses."
    ))

    add_heading(doc, "12.2 Phase 2: Expanded Coverage (March 6)", level=2)
    add_body(doc, (
        "After recalibration: 26 sports, sport-scoped abbreviations, dynamic tennis matching, "
        "exotic market filtering, 96h date window. Matched 96 markets, placed 87 trades in 2.7 hours."
    ))
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Positions opened", "87 (all on Mar 6)"],
            ["Sports covered", "6+ (EPL, Bundesliga, La Liga, NBA, NHL, ATP, WTA)"],
            ["Total exposure", "$331 of $400 max (83%)"],
            ["Avg position size", "$3.76"],
            ["Position range", "$0.53 - $8.00"],
            ["Avg edge at entry", "~6.2%"],
            ["Fill rate (steady state)", "$30/hr, 13 trades/hr"],
            ["Unrealized MTM P&L", "+$2.63 (86/87 priced)"],
            ["Resolved trades", "0 (events pending)"],
        ],
        col_aligns=["left", "center"],
    )

    add_heading(doc, "12.3 Bugs Found During Paper Trading", level=2)
    add_table(doc,
        ["Bug", "Impact", "Fix"],
        [
            ["Spread/total tolerance 1.0/0.5", "Matched wrong lines, 30-82% phantom edges",
             "Changed to 0.01 (exact match)"],
            ["No contradictory position check", "Bought Over+Under on same game",
             "Added _has_conflicting_position()"],
            ["Risk manager reset on restart", "Exceeded $300 limit, opened $694",
             "Added sync_from_tracker()"],
            ["Sport-agnostic abbreviations", "det=Pistons/Lions/Red Wings collisions",
             "Sport-scoped _SPORT_ABBREV dict"],
            ["Tennis date mismatch", "end_date=tournament end, not match date",
             "Extract date from slug for tennis"],
            ["Exotic markets as h2h/total", "First-set-winner priced vs match odds",
             "Exotic filter before type classification"],
            ["Corners-over as total", "-corners-over- matched -over- pattern",
             "Moved exotic check before totals check"],
        ],
        col_aligns=["left", "left", "left"],
    )

    add_heading(doc, "12.4 Next Steps", level=2)
    add_bullet(doc, " Wait for event resolutions (87 open positions)", bold_prefix="Immediate:")
    add_bullet(doc, " Validate CLV, win rate vs expected, Brier score at n=30+", bold_prefix="M1 (30 trades):")
    add_bullet(doc, " Binomial + t-test significance at n=100", bold_prefix="M2 (100 trades):")
    add_bullet(doc, " If CLV > +2% and flat-bet ROI > +3%, transition to live ($100 wallet)", bold_prefix="Live decision:")
    add_bullet(doc, " Esports coverage (PandaScore), more bookmakers, closing line capture", bold_prefix="Future:")

    # ── 13. Testing & Validation Framework ─────────────────────────
    add_heading(doc, "13. Testing & Validation Framework", level=1)

    add_body(doc, (
        "A comprehensive testing framework validates both code correctness and trading outcome quality. "
        "319 automated tests run in 3.7 seconds across 7 test modules."
    ))

    add_heading(doc, "13.1 Unit Test Coverage", level=2)
    add_table(doc,
        ["Module", "Tests", "Key Coverage"],
        [
            ["test_matcher.py", "142", "Team matching, tennis surnames, exotic filtering, dates, NHL/CBB abbrevs"],
            ["test_strategy.py", "24", "Edge evaluation, position sizing, conflict detection, exposure limits"],
            ["test_risk_manager.py", "26", "Exposure tracking, state persistence, limit enforcement"],
            ["test_config.py", "31", "Env loading, sport config, defaults"],
            ["test_odds_client.py", "24", "API mapping, sport keys, rate limiting"],
            ["test_polymarket_client.py", "14", "Fetch limits, market parsing"],
            ["test_trading_validation.py", "42", "Edge calibration, phantom detection, risk checks"],
            ["Total", "319", "3.7s execution time"],
        ],
        col_aligns=["left", "center", "left"],
    )

    add_heading(doc, "13.2 Trade Validator (Live Monitoring)", level=2)
    add_body(doc, (
        "The trade_validator.py module runs against live paper trading data to detect: "
        "(1) phantom edges from matching errors (exotic slug patterns, short tennis names), "
        "(2) risk limit violations (position size, exposure, contradictory positions), "
        "(3) edge calibration (Brier score, edge-bucket win rates), "
        "(4) capital efficiency (days to resolution, stale positions)."
    ))
    add_body(doc, (
        "Critical alerts are raised when: edge calibration shows systematic overestimation, "
        "phantom slug patterns account for >10% of exposure, or risk limits are breached."
    ))

    # ── References ──────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "References", level=1)

    refs = [
        "Kelly, J.L. (1956). 'A New Interpretation of Information Rate.' Bell System Technical Journal, 35(4), 917-926.",
        "Shin, H.S. (1991). 'Optimal Betting Odds Against Insider Traders.' Economic Journal, 101(408), 1179-1185.",
        "Shin, H.S. (1993). 'Measuring the Incidence of Insider Trading in a Market for State-Contingent Claims.' Economic Journal, 103(420), 1141-1153.",
        "Thorp, E.O. (2006). 'The Kelly Criterion in Blackjack, Sports Betting, and the Stock Market.' Handbook of Asset and Liability Management, Vol. 1.",
        "Clarke, S., Krase, S., Peel, D. (2017). 'Removing the Favourite-Longshot Bias.' Journal of Gambling Studies.",
        "Cheung, K. (2015). 'A Comparison of Methods for Removing the Margin from Bookmaker Odds.' Journal of Prediction Markets.",
        "Pinnacle Sports (2019). 'Closing Line Value: The Most Important Metric for Sports Bettors.'",
        "MacLean, L.C., Thorp, E.O., Ziemba, W.T. (2011). The Kelly Capital Growth Investment Criterion. World Scientific.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}]  {ref}")
        run.font.name = GARAMOND
        run.font.size = Pt(10)

    # ── Footer ──────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("EverestQuant  |  Everest Agentic AI Trader  |  March 2026  |  Internal")
    run.font.name = GARAMOND
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.italic = True

    # ── Save ────────────────────────────────────────────────────────
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    generate()
